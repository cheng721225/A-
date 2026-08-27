import io
import os
import json
import base64
from datetime import datetime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw
import streamlit as st
from streamlit_image_coordinates import streamlit_image_coordinates
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.drawing.image import Image as OpenpyxlImage

# 定義資料庫檔案名稱
DB_FILE = "defect_database.json"

# 載入資料庫的函數
def load_db():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return {}
    return {}

# 儲存資料庫的函數
def save_db(data):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Word 表格輔助函數：設定儲存格背景色
def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

# Word 表格輔助函數：設定儲存格邊框
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

st.set_page_config(
    page_title="建築工地互動式缺失查驗系統", page_icon="🏗️", layout="wide"
)

st.title("🏗️ 建築工地互動式平面圖缺失查驗系統")
st.markdown(
    "操作流程：**1. 選擇戶別與樓層** ➔ **2. 直接點擊平面圖任意位置新增缺失編號** ➔"
    " **3. 於下方表格填寫細節並上傳照片** ➔ **4. 一鍵儲存並匯出 Word / Excel 報告**"
)

# 系統啟動時，自動從本地 JSON 檔案把歷史資料撈出來
if "project_data" not in st.session_state:
    st.session_state.project_data = load_db()

st.sidebar.header("📌 1. 查驗位置設定")
unit_choice = st.sidebar.selectbox(
    "選擇戶別", ["A1戶別", "A2戶別", "A3戶別", "B1戶別", "B2戶別", "B3戶別"]
)

floor_choice = st.sidebar.selectbox(
    "選擇樓層",
    [
        "2F", "3F", "4F", "5F", "6F", "7F", "8F",
        "9F", "10F", "11F", "12F", "13F", "14F", "15F",
    ],
)

# 組合成當前唯一的空間識別 Key (例如：A1戶別_3F)
current_key = f"{unit_choice}_{floor_choice}"

# 如果這個空間還沒有在資料庫建立過，初始化它的結構
if current_key not in st.session_state.project_data:
    st.session_state.project_data[current_key] = {
        "markers": [],
        "defect_details": {},
    }
    save_db(st.session_state.project_data)

project_info = {
    "project_name": st.sidebar.text_input("工程名稱", value="弘峻草福段A區新建工程"),
    "location": f"{unit_choice} ({floor_choice})",
    "inspector": st.sidebar.text_input("查驗人員", value=""),
    "supervisor": st.sidebar.text_input("施工負責人", value=""),
}

st.sidebar.divider()
if st.sidebar.button("🗑️ 清除目前【" + current_key + "】的所有資料"):
    st.session_state.project_data[current_key] = {
        "markers": [],
        "defect_details": {},
    }
    save_db(st.session_state.project_data)
    st.rerun()

st.subheader(f"🗺️ 2. 平面圖互動標註區 ({project_info['location']})")
st.markdown("💡 **提示：直接用滑鼠點擊下方平面圖的任意位置，即可自動在該處長出紅點與編號！**")

plan_filename = f"{unit_choice.replace('戶別', '')}_plan.jpg"

if os.path.exists(plan_filename):
    original_img = Image.open(plan_filename).convert("RGB")
    orig_w, orig_h = original_img.size

    # 強制將圖片寬高縮小為 0.5 倍
    new_w, new_h = int(orig_w * 0.5), int(orig_h * 0.5)
    base_img = original_img.resize((new_w, new_h), Image.Resampling.LANCZOS)

    draw_img = base_img.copy()
    draw = ImageDraw.Draw(draw_img)

    current_markers = st.session_state.project_data[current_key]["markers"]
    current_details = st.session_state.project_data[current_key]["defect_details"]

    for m in current_markers:
        mx, my = m["x"], m["y"]
        r = 12
        draw.ellipse([mx - r, my - r, mx + r, my + r], fill="red", outline="white", width=2)
        draw.text((mx - 4, my - 7), str(m["id"]), fill="white")

    coord = streamlit_image_coordinates(draw_img, key=f"map_{current_key}")

    if coord is not None:
        click_x, click_y = coord["x"], coord["y"]
        last_marker = current_markers[-1] if current_markers else None
        if last_marker is None or last_marker["x"] != click_x or last_marker["y"] != click_y:
            next_id = max([m["id"] for m in current_markers], default=0) + 1
            current_markers.append({"id": next_id, "x": click_x, "y": click_y})
            current_details[str(next_id)] = {
                "space": "牆面",
                "space_custom": "",
                "content": "",
                "vendor": "油漆",
                "vendor_custom": "",
                "photo_b64": "",
                "remark": "",
            }
            save_db(st.session_state.project_data)
            st.rerun()

    st.markdown("---")
    c_ctrl1, c_ctrl2 = st.columns([2, 2])
    with c_ctrl1:
        st.markdown(f"**📍 目前【{current_key}】已建立缺失數：** `{len(current_markers)}` 處")
    with c_ctrl2:
        if current_markers:
            del_target = st.selectbox(
                "選擇要刪除的編號", [m["id"] for m in current_markers], key=f"del_{current_key}"
            )
            if st.button("🗑️ 刪除選定編號", type="secondary"):
                st.session_state.project_data[current_key]["markers"] = [
                    x for x in current_markers if x["id"] != del_target
                ]
                if str(del_target) in current_details:
                    del current_details[str(del_target)]
                save_db(st.session_state.project_data)
                st.rerun()

else:
    st.warning(f"⚠️ 尚未偵測到 `{unit_choice}` 的平面圖檔案 (`{plan_filename}`)。")

st.divider()

st.subheader(f"📋 3. 缺失查驗明細填報表 ({project_info['location']})")

current_markers = st.session_state.project_data[current_key]["markers"]
current_details = st.session_state.project_data[current_key]["defect_details"]

if not current_markers:
    st.info(f"👈 目前【{current_key}】尚無缺失，請直接點擊上方平面圖新增！")
else:
    with st.form(f"defect_table_form_{current_key}"):
        for m in current_markers:
            m_id = str(m["id"])
            if m_id not in current_details:
                current_details[m_id] = {
                    "space": "牆面", "space_custom": "", "content": "",
                    "vendor": "油漆", "vendor_custom": "", "photo_b64": "", "remark": ""
                }

            detail = current_details[m_id]

            st.markdown(f"#### 🔍 缺失編號：【 #{m_id} 】")
            c1, c2, c3 = st.columns(3)

            with c1:
                space_options = ["天花板", "牆面", "地板", "廚具", "衛浴", "開關插座", "給排水", "自訂"]
                curr_space = detail["space"] if detail["space"] in space_options else "自訂"
                space_sel = st.selectbox(f"缺失位置 #{m_id}", space_options, index=space_options.index(curr_space), key=f"space_{current_key}_{m_id}")
                if space_sel == "自訂":
                    detail["space_custom"] = st.text_input(f"手動輸入位置 #{m_id}", value=detail.get("space_custom", ""), key=f"space_c_{current_key}_{m_id}")
                    detail["space"] = detail["space_custom"]
                else:
                    detail["space"] = space_sel

            with c2:
                vendor_options = ["水電", "泥作", "磁磚", "油漆", "防水", "點工", "衛浴", "廚具", "自訂"]
                curr_vendor = detail["vendor"] if detail["vendor"] in vendor_options else "自訂"
                vendor_sel = st.selectbox(f"缺失廠商 #{m_id}", vendor_options, index=vendor_options.index(curr_vendor), key=f"vendor_{current_key}_{m_id}")
                if vendor_sel == "自訂":
                    detail["vendor_custom"] = st.text_input(f"手動輸入廠商 #{m_id}", value=detail.get("vendor_custom", ""), key=f"vendor_c_{current_key}_{m_id}")
                    detail["vendor"] = detail["vendor_custom"]
                else:
                    detail["vendor"] = vendor_sel

            with c3:
                photo_file = st.file_uploader(f"新增/更換照片 #{m_id}", type=["jpg", "jpeg", "png"], key=f"photo_{current_key}_{m_id}")
                if photo_file is not None:
                    try:
                        img_upload = Image.open(photo_file)
                        img_upload.thumbnail((800, 800))
                        buffered = io.BytesIO()
                        img_upload.save(buffered, format="JPEG")
                        detail["photo_b64"] = base64.b64encode(buffered.getvalue()).decode("utf-8")
                    except Exception as e:
                        st.error(f"照片處理失敗: {e}")
                
                if detail.get("photo_b64"):
                    st.success("✅ 已儲存照片 (上傳新檔可覆蓋)")

            detail["content"] = st.text_input(f"缺失內容描述 #{m_id}", value=detail["content"], key=f"content_{current_key}_{m_id}")
            detail["remark"] = st.text_input(f"備註 #{m_id}", value=detail["remark"], key=f"remark_{current_key}_{m_id}")
            st.markdown("---")

        submitted_form = st.form_submit_button(
            f"🚀 永久儲存並同時產出 Word / Excel 報告",
            type="primary",
            use_container_width=True,
        )

# 當按下儲存並產出報表按鈕時
if "submitted_form" in locals() and submitted_form:
    # 1. 寫入本地資料庫永久保存
    save_db(st.session_state.project_data)

    # 2. 開始生成專業高質感 Word 報告
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 主標題
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"【{project_info['project_name']}】室內缺失查驗紀錄表")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(44, 62, 80) # 專業深鐵灰
    title_p.paragraph_format.space_after = Pt(12)

    # 專案資訊看板 (Meta Table)
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("工程名稱", project_info["project_name"], "查驗日期", datetime.today().strftime("%Y年%m月%d日")),
        ("施作位置", project_info["location"], "查驗人員 / 主管", f"{project_info['inspector']} / {project_info['supervisor']}"),
    ]
    
    meta_col_widths = [Inches(1.2), Inches(2.3), Inches(1.2), Inches(2.3)]
    for r_idx, row_content in enumerate(meta_data):
        for c_idx, text_val in enumerate(row_content):
            cell = meta_table.cell(r_idx, c_idx)
            cell.text = str(text_val)
            cell.width = meta_col_widths[c_idx]
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            set_cell_background(cell, "F8F9F9" if r_idx % 2 == 0 else "F2F4F4")
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Arial"
                if c_idx % 2 == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(85, 85, 85)
                else:
                    run.font.color.rgb = RGBColor(34, 34, 34)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 一、 平面圖區塊
    if os.path.exists(plan_filename):
        h1_p = doc.add_paragraph()
        h1_run = h1_p.add_run(f"一、 戶別平面圖與標註位置 ({project_info['location']})")
        h1_run.bold = True
        h1_run.font.size = Pt(13)
        h1_run.font.color.rgb = RGBColor(52, 73, 94)
        h1_p.paragraph_format.space_before = Pt(10)
        h1_p.paragraph_format.space_after = Pt(6)

        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_byte_arr = io.BytesIO()
        draw_img.save(img_byte_arr, format="JPEG")
        img_byte_arr.seek(0)
        img_p.add_run().add_picture(img_byte_arr, width=Inches(4.8))
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 二、 缺失查驗明細表區塊
    h2_p = doc.add_paragraph()
    h2_run = h2_p.add_run("二、 缺失查驗明細表")
    h2_run.bold = True
    h2_run.font.size = Pt(13)
    h2_run.font.color.rgb = RGBColor(52, 73, 94)
    h2_p.paragraph_format.space_before = Pt(10)
    h2_p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["編號", "缺失位置", "缺失內容", "缺失廠商", "現場照片", "備註"]
    col_widths = [Inches(0.6), Inches(1.1), Inches(2.0), Inches(1.1), Inches(1.5), Inches(1.2)]
    
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "2C3E50") # 深鐵灰表頭
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=100, right=100)
        
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(255, 255, 255)

    for idx, m in enumerate(current_markers):
        m_id = str(m["id"])
        detail = current_details.get(m_id, {})
        row_cells = table.add_row().cells

        row_values = [
            f"#{m_id}",
            detail.get("space", ""),
            detail.get("content", ""),
            detail.get("vendor", ""),
            "", # 照片後續放入
            detail.get("remark", "")
        ]

        for i, val in enumerate(row_values):
            row_cells[i].text = val
            row_cells[i].width = col_widths[i]
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            if idx % 2 == 1:
                set_cell_background(row_cells[i], "F9F9F9") # 斑馬紋

            p = row_cells[i].paragraphs[0]
            if i in [0]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = "Arial"
                run.font.color.rgb = RGBColor(51, 51, 51)

        # 放入現場照片
        photo_b64 = detail.get("photo_b64")
        if photo_b64:
            try:
                photo_para = row_cells[4].paragraphs[0]
                photo_para.text = "" # 清空文字
                photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = photo_para.add_run()
                img_bytes = base64.b64decode(photo_b64)
                run.add_picture(io.BytesIO(img_bytes), width=Inches(1.2))
                row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except:
                row_cells[4].text = "載入失敗"
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            row_cells[4].text = "無照片"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    word_io = io.BytesIO()
    doc.save(word_io)
    word_io.seek(0)

    # 3. 開始生成專業高質感 Excel 報表
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "缺失查驗明細表"
    ws.views.sheetView[0].showGridLines = True

    font_title = Font(name="Arial", size=16, bold=True, color="2C3E50")
    font_section = Font(name="Arial", size=13, bold=True, color="34495E")
    font_meta_val = Font(name="Arial", size=10, color="222222")
    font_header = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=10, color="333333")

    fill_header = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    fill_meta = PatternFill(start_color="F8F9F9", end_color="F8F9F9", fill_type="solid")
    fill_zebra = PatternFill(start_color="F2F4F4", end_color="F2F4F4", fill_type="solid")

    border_thin = Border(
        left=Side(style='thin', color='D5D8DC'), right=Side(style='thin', color='D5D8DC'),
        top=Side(style='thin', color='D5D8DC'), bottom=Side(style='thin', color='D5D8DC')
    )
    border_meta = Border(
        left=Side(style='thin', color='ABB2B9'), right=Side(style='thin', color='ABB2B9'),
        top=Side(style='thin', color='ABB2B9'), bottom=Side(style='thin', color='ABB2B9')
    )

    align_center = Alignment(horizontal='center', vertical='center')
    align_left = Alignment(horizontal='left', vertical='center', wrap_text=True)

    ws.append([f"【{project_info['project_name']}】室內缺失查驗紀錄表"])
    ws.cell(row=1, column=1).font = font_title
    ws.row_dimensions[1].height = 30
    
    ws.append([f"工程名稱: {project_info['project_name']}", "", f"查驗日期: {datetime.today().strftime('%Y年%m月%d日')}"])
    ws.append([f"施作位置: {project_info['location']}", "", f"查驗人員 / 主管: {project_info['inspector']} / {project_info['supervisor']}"])
    
    ws.row_dimensions[2].height = 20
    ws.row_dimensions[3].height = 20

    for r in [2, 3]:
        for c in range(1, 7):
            cell = ws.cell(row=r, column=c)
            cell.fill = fill_meta
            cell.border = border_meta
            cell.font = font_meta_val

    ws.append([])
    ws.row_dimensions[4].height = 15

    ws.append([f"一、 戶別平面圖與標註位置 ({project_info['location']})"])
    ws.cell(row=5, column=1).font = font_section
    ws.row_dimensions[5].height = 25
    ws.append([])

    if os.path.exists(plan_filename):
        plan_img_io = io.BytesIO()
        draw_img.save(plan_img_io, format="JPEG")
        plan_img_io.seek(0)
        
        excel_plan_img = OpenpyxlImage(plan_img_io)
        excel_plan_img.width = 380  
        excel_plan_img.height = int(380 * (draw_img.height / draw_img.width))
        ws.add_image(excel_plan_img, "A7") 

    for _ in range(16):
        ws.append([])

    sec_row = ws.max_row + 1
    ws.append(["二、 缺失查驗明細表"])
    ws.cell(row=sec_row, column=1).font = font_section
    ws.row_dimensions[sec_row].height = 25

    excel_headers = ["編號", "缺失位置", "缺失內容", "缺失廠商", "現場照片", "備註"]
    ws.append(excel_headers)
    table_header_row_idx = ws.max_row
    ws.row_dimensions[table_header_row_idx].height = 26

    for col_num in range(1, 7):
        cell = ws.cell(row=table_header_row_idx, column=col_num)
        cell.fill = fill_header
        cell.font = font_header
        cell.alignment = align_center
        cell.border = border_thin

    ws.column_dimensions['A'].width = 10
    ws.column_dimensions['B'].width = 16
    ws.column_dimensions['C'].width = 32
    ws.column_dimensions['D'].width = 16
    ws.column_dimensions['E'].width = 18
    ws.column_dimensions['F'].width = 32

    start_row_for_details = table_header_row_idx + 1

    for idx, m in enumerate(current_markers):
        current_row = start_row_for_details + idx
        ws.row_dimensions[current_row].height = 80

        m_id = str(m["id"])
        detail = current_details.get(m_id, {})
        
        row_values = [
            f"#{m_id}",
            detail.get("space", ""),
            detail.get("content", ""),
            detail.get("vendor", ""),
            "",
            detail.get("remark", "")
        ]
        
        for col_num, val in enumerate(row_values, 1):
            cell = ws.cell(row=current_row, column=col_num, value=val)
            cell.font = font_body
            cell.border = border_thin
            
            if idx % 2 == 1:
                cell.fill = fill_zebra

            if col_num in [1, 5]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

        photo_b64 = detail.get("photo_b64")
        if photo_b64:
            try:
                img_bytes = base64.b64decode(photo_b64)
                img_io = io.BytesIO(img_bytes)
                
                excel_cell_img = OpenpyxlImage(img_io)
                excel_cell_img.width = 75
                excel_cell_img.height = 75
                
                cell_coordinate = f"E{current_row}"
                ws.add_image(excel_cell_img, cell_coordinate)
            except:
                ws.cell(row=current_row, column=5, value="載入失敗").alignment = align_center
        else:
            ws.cell(row=current_row, column=5, value="無照片").alignment = align_center

    excel_io = io.BytesIO()
    wb.save(excel_io)
    excel_io.seek(0)

    st.success("🎉 資料已永久儲存，並成功產出專業高質感版型的 Word 與 Excel 報告！請點擊下方按鈕下載：")
    
    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        st.download_button(
            label=f"📥 下載 Word 報告 (.docx)",
            data=word_io,
            file_name=f"室內缺失查驗紀錄_{project_info['location'].replace(' ', '_')}_{datetime.today().strftime('%Y%m%d')}.docx",
            mime=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            use_container_width=True
        )
    with col_dl2:
        st.download_button(
            label=f"📊 下載 Excel 報表 (.xlsx)",
            data=excel_io,
            file_name=f"室內缺失查驗明細_{project_info['location'].replace(' ', '_')}_{datetime.today().strftime('%Y%m%d')}.xlsx",
            mime=("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
            use_container_width=True
        )
